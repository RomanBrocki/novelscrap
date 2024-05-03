def copiarcapitulos(site):
    from selenium import webdriver
    import time
    from docx import Document
    from selenium.webdriver.common.by import By
    navegador = webdriver.Chrome()  
    while site != "":
        navegador.get(site)
        time.sleep(1.5)
        titulo = navegador.find_element(By.CLASS_NAME, 'chr-text')
        conteudo = navegador.find_element(By.ID, 'chr-content')
        documento = Document("Novel.docx")
        paragrafo1 = documento.add_paragraph(titulo.text)
        paragrafo2 = documento.add_paragraph(conteudo.text)
        documento.add_page_break()
        documento.save("Novel.docx")
        links = navegador.find_elements(By.TAG_NAME, 'a')
        listalinks = []
        for lnk in links:
            listalinks.append(lnk.get_attribute('href'))
        site = listalinks[-8]